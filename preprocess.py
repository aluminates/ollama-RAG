import os
import csv
import warnings
import pdfplumber
import tabula
from docx import Document as DocxDocument
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.document_loaders import UnstructuredPowerPointLoader
from dotenv import load_dotenv

load_dotenv()
warnings.filterwarnings("ignore")

def load_csv_documents(file_path):
    documents = []
    with open(file_path, mode='r', encoding='utf-8') as file:
        reader = csv.reader(file)
        header = next(reader)
        for row in reader:
            content = ' '.join(row)
            documents.append(Document(page_content=content, metadata={"source": file_path, "type": "csv"}))
    return documents

class PPTExtraction:
    def __init__(self, file_path):
        self.file_path = file_path
        self.loader = UnstructuredPowerPointLoader(self.file_path, mode="elements")
        self.data = self.loader.load()

    def extract(self):
        slides = []
        current_slide_number = None

        for document in self.data:
            if document.metadata["category"] == "Title":
                slide_number = document.metadata["page_number"]
                if slide_number != current_slide_number:
                    if slide_number == 1:
                        slide = f"Slide {slide_number}:\n\nTitle: {document.page_content}"
                    else:
                        slide = f"Slide {slide_number}:\n\nOutline: {document.page_content}"
                    current_slide_number = slide_number
                else:
                    slide = f"Outline: {document.page_content}"
            elif document.metadata["category"] in ["NarrativeText", "ListItem"]:
                slide = f"Content: {document.page_content}"
            elif document.metadata["category"] == "PageBreak":
                slide = ""
                current_slide_number = None
            else:
                continue

            slides.append(slide)

        formatted_slides = "\n\n".join(slides)
        return Document(page_content=formatted_slides, metadata={"source": self.file_path, "type": "pptx"})

def load_docx_documents(file_path):
    documents = []
    doc = DocxDocument(file_path)
    content = []

    def process_paragraph(paragraph):
        if paragraph.style.name.startswith('Heading'):
            return f"{paragraph.style.name}: {paragraph.text}\n"
        elif paragraph.style.name == 'List Paragraph':
            return f"- {paragraph.text}\n"
        else:
            return f"{paragraph.text}\n"

    def process_table(table):
        table_content = []
        for row in table.rows:
            row_content = [cell.text for cell in row.cells]
            table_content.append(" | ".join(row_content))
        return "Table:\n" + "\n".join(table_content) + "\n"

    for paragraph in doc.paragraphs:
        content.append(process_paragraph(paragraph))
    
    for table in doc.tables:
        content.append(process_table(table))

    full_content = "".join(content)
    documents.append(Document(page_content=full_content, metadata={"source": file_path, "type": "docx"}))
    return documents

def load_pdf_documents(file_path):
    documents = []
    
    # Extract text
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                documents.append(Document(page_content=text, metadata={"source": f"{file_path} (page {page_num + 1})", "type": "pdf"}))
    
    # Extract tables
    tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True)
    for i, table in enumerate(tables):
        table_content = table.to_string(index=False)
        documents.append(Document(page_content=table_content, metadata={"source": f"{file_path} (table {i + 1})", "type": "pdf_table"}))
    return documents

def load_documents(data_path):
    documents = []

    for file in os.listdir(data_path):
        file_path = os.path.join(data_path, file)
        if file.endswith(".csv"):
            csv_documents = load_csv_documents(file_path)
            documents.extend(csv_documents)
        elif file.endswith(".pptx"):
            ppt_extraction = PPTExtraction(file_path)
            documents.append(ppt_extraction.extract())
        elif file.endswith(".docx"):
            docx_documents = load_docx_documents(file_path)
            documents.extend(docx_documents)
        elif file.endswith(".pdf"):
            pdf_documents = load_pdf_documents(file_path)
            documents.extend(pdf_documents)
    
    return documents

def hierarchical_split(documents):
    doc_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    content_splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=40)
    
    split_docs = []
    for doc in documents:
        doc_chunks = doc_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(doc_chunks):
            sub_chunks = content_splitter.split_text(chunk)
            for j, sub_chunk in enumerate(sub_chunks):
                split_docs.append(Document(
                    page_content=sub_chunk,
                    metadata={
                        **doc.metadata,
                        "chunk": i,
                        "sub_chunk": j
                    }
                ))
    return split_docs

def create_vector_db():
    data_path = os.getenv('DATA_PATH')
    db_path = os.getenv('DB_PATH')

    if not data_path or not db_path:
        raise ValueError("DATA_PATH or DB_PATH environment variables not set.")

    documents = load_documents(data_path)
    print(f"Processed {len(documents)} documents.")

    texts = hierarchical_split(documents)
    print(f"Split into {len(texts)} chunks.")

    vector_store = Chroma.from_documents(
        documents=texts,
        embedding=HuggingFaceEmbeddings(model_name='sentence-transformers/all-mpnet-base-v2'),
        persist_directory=db_path
    )
    vector_store.persist()
    print(f"Vector database persisted at {db_path}.")

if __name__ == "__main__":
    create_vector_db()
